#!/usr/bin/env python3
"""
Generate resume_sections.tex from resume_content.yaml.

Reads Summary, Skills, Experience, and Projects from the YAML content file,
escapes LaTeX special characters, and renders the four sections in Awesome-CV format.
Supports optional AI tailoring (--tailor) to adapt content to a job description.

Copyright (C) 2025  Manoj Ponagandla

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with this program.  If not, see <https://www.gnu.org/licenses/>.
"""

import argparse
import logging
import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

try:
    import yaml
except ImportError:
    print("Error: PyYAML is required. Install with: pip install pyyaml", file=sys.stderr)
    sys.exit(1)

__version__ = "1.0.0"

PROJECT_ROOT = Path(__file__).resolve().parent
RESOURCES_DIR = PROJECT_ROOT / "resources"
CONTENT_FILE = PROJECT_ROOT / "my-content" / "resume_content.yaml"
PROFILE_FILE = PROJECT_ROOT / "my-content" / "user_profile.md"
OUTPUT_FILE = RESOURCES_DIR / "resume_sections.tex"

# Max characters per skill items line so PDF stays one line (Awesome-CV ~70% text width).
SKILL_ITEMS_MAX_CHARS = 58


def escape_latex(text: str) -> str:
    """Escape LaTeX special characters in plain text."""
    if not text:
        return ""
    replacements = [
        ("\\", r"\textbackslash "),
        ("&", r"\&"),
        ("%", r"\%"),
        ("#", r"\#"),
        ("_", r"\_"),
        ("{", r"\{"),
        ("}", r"\}"),
        ("$", r"\$"),
        ("~", r"\textasciitilde "),
        ("^", r"\textasciicircum "),
    ]
    result = str(text)
    for old, new in replacements:
        result = result.replace(old, new)
    return result


def render_summary(summary: str) -> str:
    """Render the Summary section."""
    escaped = escape_latex(summary.strip())
    return f"""%--------------------------------------------------
% Summary
%--------------------------------------------------
\\cvsection{{Summary}}

\\begin{{cvparagraph}}
{escaped}
\\end{{cvparagraph}}
"""


def _truncate_skill_items(items_str: str, max_chars: int) -> str:
    """Limit items string to max_chars so it fits on one line; truncate at last comma if over."""
    if not items_str or len(items_str) <= max_chars:
        return items_str
    candidate = items_str[: max_chars + 1]
    last_comma = candidate.rfind(",")
    if last_comma > 0:
        return items_str[:last_comma].strip()
    return items_str[:max_chars].strip()


def render_skills(skills: list[dict]) -> str:
    """Render the Skills section. One row per category; items truncated to fit one line."""
    lines = [
        "%--------------------------------------------------",
        "% Skills (condensed, Awesome-CV style)",
        "%--------------------------------------------------",
        "\\cvsection{Skills}",
        "",
        "\\begin{cvskills}",
    ]
    for skill in skills:
        category = escape_latex(skill["category"])
        raw_items = skill.get("items", "")
        # Schema expects items as string; LLM may output a list — normalize to one string.
        if isinstance(raw_items, (list, tuple)):
            items_str = ", ".join(str(x).strip() for x in raw_items)
        else:
            items_str = str(raw_items).strip() if raw_items else ""
        items_str = _truncate_skill_items(items_str, SKILL_ITEMS_MAX_CHARS)
        items = escape_latex(items_str)
        items = items.replace(", ", ", \\allowbreak ")
        lines.append(f"\\cvskill{{{category}}}{{{items}}}")
    lines.extend(["\\end{cvskills}", ""])
    return "\n".join(lines)


def render_entry(
    position: str,
    organization: str,
    date: str,
    location: str,
    bullets: list[str],
    *,
    raw_position: bool = False,
) -> str:
    """Render a single cventry (experience or project)."""
    pos_text = position if raw_position else escape_latex(position)
    org_text = escape_latex(organization)
    date_text = escape_latex(date)
    loc_text = escape_latex(location)

    # cventry: position, title, location, date, description
    bullet_lines = [f"\\item {escape_latex(b)}" for b in bullets]
    items_block = "\\begin{cvitems}\n" + "\n".join(bullet_lines) + "\n\\end{cvitems}"
    return f"""\\cventry
{{{pos_text}}}
{{{org_text}}}
{{{date_text}}}
{{{loc_text}}}
{{
{items_block}
}}
"""


def render_experience(experience: list[dict]) -> str:
    """Render the Experience section."""
    lines = [
        "%--------------------------------------------------",
        "% Experience",
        "%--------------------------------------------------",
        "\\cvsection{Experience}",
        "",
    ]
    for entry in experience:
        lines.append(
            render_entry(
                position=entry["position"],
                organization=entry.get("organization", ""),
                date=entry.get("date", ""),
                location=entry.get("location", ""),
                bullets=entry.get("bullets", []),
                raw_position=entry.get("raw_position", False),
            )
        )
    return "\n".join(lines)


def render_projects(projects: list[dict]) -> str:
    """Render the Projects section."""
    lines = [
        "%--------------------------------------------------",
        "% Projects",
        "%--------------------------------------------------",
        "\\cvsection{Projects}",
        "",
    ]
    for entry in projects:
        # LLM may output "name" instead of "position" for projects; accept both.
        position = entry.get("position") or entry.get("name") or ""
        bullets = entry.get("bullets")
        if bullets is None and "description" in entry:
            d = entry["description"]
            bullets = [d] if isinstance(d, str) else list(d) if isinstance(d, (list, tuple)) else []
        if bullets is None:
            bullets = []
        lines.append(
            render_entry(
                position=position,
                organization=entry.get("organization", ""),
                date=entry.get("date", ""),
                location=entry.get("location", ""),
                bullets=bullets,
                raw_position=entry.get("raw_position", False),
            )
        )
    return "\n".join(lines)


def compile_latex_to_pdf(resources_dir: Path | None = None) -> Path | None:
    """
    Compile resume.tex to PDF using xelatex. Returns the path to the generated PDF,
    or None if compilation fails (e.g. xelatex not installed).
    """
    resources_dir = resources_dir or RESOURCES_DIR
    resume_tex = resources_dir / "resume.tex"
    if not resume_tex.exists():
        return None
    try:
        subprocess.run(
            ["xelatex", "-interaction=batchmode", "resume.tex"],
            cwd=resources_dir,
            check=True,
            capture_output=True,
        )
    except (subprocess.CalledProcessError, FileNotFoundError):
        return None
    pdf_path = resources_dir / "resume.pdf"
    if not pdf_path.exists():
        return None
    # Copy to output/ with timestamp (like Makefile)
    output_dir = PROJECT_ROOT / "output"
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    output_pdf = output_dir / f"resume-{timestamp}.pdf"
    shutil.copy2(pdf_path, output_pdf)
    return output_pdf


def render_yaml_to_latex(yaml_str_or_data: str | dict, output_path: Path) -> None:
    """
    Render resume YAML (string or dict) to LaTeX and write to output_path.
    Used by main() and by the chat agent's generate_resume tool.
    """
    if isinstance(yaml_str_or_data, str):
        data = yaml.safe_load(yaml_str_or_data)
    else:
        data = yaml_str_or_data
    if not data:
        raise ValueError("Content is empty.")
    sections = []
    if "summary" in data:
        sections.append(render_summary(data["summary"]))
    if "skills" in data:
        sections.append(render_skills(data["skills"]))
    if "experience" in data:
        sections.append(render_experience(data["experience"]))
    if "projects" in data:
        sections.append(render_projects(data["projects"]))
    output = "\n".join(sections)
    output_path = output_path.resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    gen_time = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(f"% Generated: {gen_time}\n")
        f.write(output)


def _extract_header_from_resume_tex(resume_tex_path: Path) -> tuple[str, str]:
    """Extract name and contact line from resume.tex. Returns (full_name, contact_line)."""
    if not resume_tex_path.exists():
        return ("", "")
    text = resume_tex_path.read_text(encoding="utf-8")
    first, last = "", ""
    m = re.search(r"\\name\{([^{}]*)\}\{([^{}]*)\}", text)
    if m:
        first, last = m.group(1).strip(), m.group(2).strip()
    full_name = f"{first} {last}".strip() if (first or last) else ""

    contact = ""
    m = re.search(r"\\position\{(.*?)\}", text, re.DOTALL)
    if m:
        raw = m.group(1)
        raw = re.sub(r"\\small\s*", "", raw)
        raw = re.sub(r"\\href\{[^{}]*\}\{([^{}]*)\}", r"\1", raw)
        raw = re.sub(r"\s+", " ", raw).strip()
        contact = raw

    return (full_name, contact)


def _load_personal_info_from_config() -> tuple[str, str] | None:
    """Load personal_info from resume_config.yaml if present. Returns (full_name, contact) or None."""
    config_path = PROJECT_ROOT / "resume_config.yaml"
    if not config_path.exists():
        return None
    try:
        with open(config_path, encoding="utf-8") as f:
            config = yaml.safe_load(f)
        if not config or "personal_info" not in config:
            return None
        pi = config["personal_info"]
        first = pi.get("first_name", "")
        last = pi.get("last_name", "")
        full_name = f"{first} {last}".strip() if (first or last) else ""
        contact = pi.get("contact", "")
        return (full_name, contact)
    except Exception:
        return None


def render_yaml_to_docx(yaml_str_or_data: str | dict, output_path: Path) -> Path:
    """
    Render resume YAML to a Word document matching the Awesome-CV PDF formatting.
    Returns the path to the generated .docx file.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError:
        raise ImportError("python-docx is required for Word output. Install with: pip install python-docx")

    if isinstance(yaml_str_or_data, str):
        data = yaml.safe_load(yaml_str_or_data)
    else:
        data = yaml_str_or_data
    if not data:
        raise ValueError("Content is empty.")

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.5)

    font_name = "Source Sans Pro"

    def _remove_table_borders(table) -> None:
        """Remove all borders from a table for a clean, borderless look."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is None:
            tblBorders = OxmlElement("w:tblBorders")
            tblPr.append(tblBorders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = f"w:{edge}"
            el = tblBorders.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                tblBorders.append(el)
            el.set(qn("w:val"), "nil")
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn("w:tcBorders"))
                if tcBorders is None:
                    tcBorders = OxmlElement("w:tcBorders")
                    tcPr.append(tcBorders)
                for edge in ("top", "left", "bottom", "right"):
                    tag = f"w:{edge}"
                    el = tcBorders.find(qn(tag))
                    if el is None:
                        el = OxmlElement(tag)
                        tcBorders.append(el)
                    el.set(qn("w:val"), "nil")

    # Header: name + contact
    personal = _load_personal_info_from_config()
    if personal and (personal[0] or personal[1]):
        full_name, contact = personal
    else:
        full_name, contact = _extract_header_from_resume_tex(RESOURCES_DIR / "resume.tex")

    if full_name:
        p = doc.add_paragraph()
        run = p.add_run(full_name)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = font_name
        p.paragraph_format.space_after = Pt(1)
    if contact:
        p = doc.add_paragraph()
        run = p.add_run(contact)
        run.font.size = Pt(8)
        run.font.name = font_name
        run.font.all_caps = True
        p.paragraph_format.space_after = Pt(2)

    def add_section_heading(title: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = font_name
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.border_bottom_width = Pt(0.5)
        p.paragraph_format.space_after = Pt(1)

    def add_summary_paragraph(text: str) -> None:
        p = doc.add_paragraph(text.strip())
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.name = font_name
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

    def add_skills_table(skills: list[dict]) -> None:
        table = doc.add_table(rows=len(skills), cols=2)
        _remove_table_borders(table)
        table.autofit = False
        # Category column: enough space to avoid wrapping; items column takes the rest
        cat_width = Cm(4.5)
        items_width = Cm(14.0)
        for row in table.rows:
            row.cells[0].width = cat_width
            row.cells[1].width = items_width
        for i, skill in enumerate(skills):
            cat = skill.get("category", "")
            raw_items = skill.get("items", "")
            if isinstance(raw_items, (list, tuple)):
                items_str = ", ".join(str(x).strip() for x in raw_items)
            else:
                items_str = str(raw_items).strip() if raw_items else ""
            row = table.rows[i]
            row.cells[0].text = cat
            row.cells[1].text = items_str
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.name = font_name
            if row.cells[0].paragraphs[0].runs:
                row.cells[0].paragraphs[0].runs[0].bold = True
                row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            if row.cells[1].paragraphs[0].runs:
                row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)

    def add_entry(position: str, organization: str, date: str, location: str, bullets: list[str]) -> None:
        has_org_loc = bool(organization or location)
        table = doc.add_table(rows=2 if has_org_loc else 1, cols=2)
        _remove_table_borders(table)
        table.autofit = False
        # Left column (org/position) wider, right column (date/location) narrower - matches PDF
        for row in table.rows:
            row.cells[0].width = Cm(14.0)
            row.cells[1].width = Cm(4.5)
        if has_org_loc:
            table.rows[0].cells[0].text = organization or ""
            table.rows[0].cells[1].text = location or ""
            table.rows[1].cells[0].text = position
            table.rows[1].cells[1].text = date or ""
        else:
            table.rows[0].cells[0].text = position
            table.rows[0].cells[1].text = date or ""
        def style_run(cell_idx: int, row_idx: int, **kwargs) -> None:
            cell = table.rows[row_idx].cells[cell_idx]
            if cell.paragraphs and cell.paragraphs[0].runs:
                r = cell.paragraphs[0].runs[0]
                for k, v in kwargs.items():
                    setattr(r.font, k, v)

        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = font_name
        # Row 0: org/position (bold 10pt) | location (italic 9pt)
        style_run(0, 0, bold=True, size=Pt(10))
        style_run(1, 0, italic=True, size=Pt(9))
        # Row 1 (when has_org_loc): position (8pt) | date (italic 8pt)
        if has_org_loc:
            style_run(0, 1, size=Pt(8))
            style_run(1, 1, italic=True, size=Pt(8))
        for b in bullets:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(b)
            run.font.size = Pt(9)
            run.font.name = font_name
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)

    if "summary" in data:
        add_section_heading("Summary")
        add_summary_paragraph(data["summary"])

    if "skills" in data:
        add_section_heading("Skills")
        add_skills_table(data["skills"])

    if "experience" in data:
        add_section_heading("Experience")
        for entry in data["experience"]:
            add_entry(
                position=entry.get("position", ""),
                organization=entry.get("organization", ""),
                date=entry.get("date", ""),
                location=entry.get("location", ""),
                bullets=entry.get("bullets", []),
            )

    if "projects" in data:
        add_section_heading("Projects")
        for entry in data["projects"]:
            position = entry.get("position") or entry.get("name") or ""
            bullets = entry.get("bullets")
            if bullets is None and "description" in entry:
                d = entry["description"]
                bullets = [d] if isinstance(d, str) else list(d) if isinstance(d, (list, tuple)) else []
            if bullets is None:
                bullets = []
            add_entry(
                position=position,
                organization=entry.get("organization", ""),
                date=entry.get("date", ""),
                location=entry.get("location", ""),
                bullets=bullets,
            )

    output_path = output_path.resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate resume_sections.tex from resume content YAML. Optionally tailor content to a job description using an LLM."
    )
    parser.add_argument(
        "-i", "--input",
        type=Path,
        default=CONTENT_FILE,
        help="Path to content YAML (default: my-content/resume_content.yaml)",
    )
    parser.add_argument(
        "-o", "--output",
        type=Path,
        default=OUTPUT_FILE,
        help="Path for generated LaTeX sections (default: resources/resume_sections.tex)",
    )
    parser.add_argument(
        "--tailor",
        metavar="PATH",
        help="Path to job description file (text or YAML with 'description' key). Tailors content via LLM.",
    )
    parser.add_argument(
        "--tailor-url",
        metavar="URL",
        help="URL of job description to fetch and tailor content to.",
    )
    parser.add_argument(
        "--no-tailor",
        action="store_true",
        help="Disable LLM tailoring even if --tailor/--tailor-url would be inferred.",
    )
    parser.add_argument(
        "--openai",
        action="store_true",
        help="Use OpenAI API for tailoring (default: Ollama). Requires OPENAI_API_KEY.",
    )
    parser.add_argument(
        "--version",
        action="version",
        version=f"%(prog)s {__version__}",
    )
    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Enable verbose (debug) output.",
    )
    parser.add_argument(
        "--max-fact-error-rate",
        type=float,
        default=None,
        metavar="RATE",
        help="Max allowed share of LLM-introduced facts (0.0-1.0). Overrides config/env. Within limit, tailored content is accepted; over limit triggers a targeted rewrite of offending entries.",
    )
    parser.add_argument(
        "-f", "--format",
        choices=["pdf", "docx", "both"],
        default="pdf",
        help="Output format: pdf (default), docx (Word), or both.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.verbose:
        logging.basicConfig(level=logging.DEBUG)

    content_path = args.input.resolve()
    if not content_path.exists():
        print(f"Error: Content file not found: {content_path}", file=sys.stderr)
        sys.exit(1)

    tailor_source = None
    if not args.no_tailor and (args.tailor or args.tailor_url):
        tailor_source = args.tailor or args.tailor_url
        if args.tailor:
            jd_path = Path(args.tailor).resolve()
            if not jd_path.exists():
                print(f"Error: Job description file not found: {jd_path}", file=sys.stderr)
                sys.exit(1)
            tailor_source = str(jd_path)

    if tailor_source:
        import tailor as tailor_mod
        profile_path = PROFILE_FILE.resolve()
        if profile_path.exists():
            print("Tailoring resume from user profile + job description via LLM...", flush=True)
            yaml_str = tailor_mod.tailor_from_profile(
                profile_path,
                tailor_source,
                use_openai=args.openai,
                verbose=args.verbose,
                max_fact_error_rate=args.max_fact_error_rate,
            )
            if yaml_str is not None:
                data = yaml.safe_load(yaml_str)
            else:
                # Fall back to resume_content.yaml and YAML-based tailoring
                print("Falling back to resume_content.yaml for tailoring.", flush=True)
                yaml_str = tailor_mod.tailor(
                    content_path,
                    tailor_source,
                    use_openai=args.openai,
                    verbose=args.verbose,
                    max_fact_error_rate=args.max_fact_error_rate,
                )
                data = yaml.safe_load(yaml_str)
        else:
            print("Tailoring resume to job description via LLM (Ollama)...", flush=True)
            yaml_str = tailor_mod.tailor(
                content_path,
                tailor_source,
                use_openai=args.openai,
                verbose=args.verbose,
                max_fact_error_rate=args.max_fact_error_rate,
            )
            data = yaml.safe_load(yaml_str)
    else:
        with open(content_path, encoding="utf-8") as f:
            data = yaml.safe_load(f)

    if not data:
        print("Error: Content file is empty.", file=sys.stderr)
        sys.exit(1)

    output_dir = PROJECT_ROOT / "output"
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")

    if args.format in ("pdf", "both"):
        out_path = args.output.resolve()
        render_yaml_to_latex(data, out_path)
        print(f"Generated {out_path}")
        if args.format == "pdf":
            pdf_path = compile_latex_to_pdf()
            if pdf_path:
                print(f"Saved {pdf_path}")
        else:
            pdf_path = compile_latex_to_pdf()
            if pdf_path:
                print(f"Saved {pdf_path}")

    if args.format in ("docx", "both"):
        docx_path = output_dir / f"resume-{timestamp}.docx"
        render_yaml_to_docx(data, docx_path)
        print(f"Generated {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "chat":
        sys.argv = [sys.argv[0]] + sys.argv[2:]
        from resume_agent import main as chat_main
        chat_main()
    else:
        main()
